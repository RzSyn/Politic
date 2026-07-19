# generate_word.py — สร้างไฟล์ Word (.docx) จากรัฐธรรมนูญจำลองที่ประมวลผลแล้ว
# ตั้งค่าภาษาไทย + ฟอนต์ TH Sarabun New ไม่ให้ขึ้นเส้นแดง (spell check)
# หมวดใหม่ทุกหมวดจะขึ้นหน้ากระดาษใหม่เสมอ

import os
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

wd = r"c:\Users\ASUA\OneDrive\Documents\รัฐธรรมนุญจำลอง"
processed_text_path = os.path.join(wd, "files", "processed_constitution.txt")
# Handle file locking
def get_writable_path(base_path):
    if not os.path.exists(base_path):
        return base_path
    try:
        with open(base_path, 'ab') as f:
            pass
        return base_path
    except PermissionError:
        name, ext = os.path.splitext(base_path)
        for v in range(1, 100):
            vpath = f"{name}_v{v}{ext}"
            if not os.path.exists(vpath):
                return vpath
            try:
                with open(vpath, 'ab') as f:
                    pass
                return vpath
            except PermissionError:
                continue
        return f"{name}_new{ext}"

word_output_path = get_writable_path(os.path.join(wd, "files", "รัฐธรรมนูญ_๒๕๒๕_แก้ไขเพิ่มเติม.docx"))
toc_output_path = get_writable_path(os.path.join(wd, "files", "สารบัญรัฐธรรมนูญ.docx"))

# ==================== Helper Functions ====================

THAI_FONT = "TH Sarabun New"

def to_thai_numeral(num):
    thai_nums = {'0': '๐', '1': '๑', '2': '๒', '3': '๓', '4': '๔',
                 '5': '๕', '6': '๖', '7': '๗', '8': '๘', '9': '๙'}
    return ''.join(thai_nums.get(d, d) for d in str(num))

def set_thai_language(run):
    """Set the proofing language of a run to Thai (th-TH)."""
    rPr = run._element.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'th-TH')
    lang.set(qn('w:bidi'), 'th-TH')
    lang.set(qn('w:eastAsia'), 'th-TH')
    rPr.append(lang)

def set_document_thai_language(doc):
    """Set the entire document's default language to Thai."""
    style = doc.styles['Normal']
    rPr = style.element.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'th-TH')
    lang.set(qn('w:bidi'), 'th-TH')
    lang.set(qn('w:eastAsia'), 'th-TH')
    rPr.append(lang)
    
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = OxmlElement('w:docDefaults')
        styles_element.insert(0, docDefaults)
    
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault')
        docDefaults.append(rPrDefault)
    
    rPr2 = rPrDefault.find(qn('w:rPr'))
    if rPr2 is None:
        rPr2 = OxmlElement('w:rPr')
        rPrDefault.append(rPr2)
    
    lang2 = OxmlElement('w:lang')
    lang2.set(qn('w:val'), 'th-TH')
    lang2.set(qn('w:bidi'), 'th-TH')
    lang2.set(qn('w:eastAsia'), 'th-TH')
    rPr2.append(lang2)
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), THAI_FONT)
    rFonts.set(qn('w:hAnsi'), THAI_FONT)
    rFonts.set(qn('w:cs'), THAI_FONT)
    rFonts.set(qn('w:eastAsia'), THAI_FONT)
    rPr2.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '32')
    rPr2.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '32')
    rPr2.append(szCs)

def add_thai_run(paragraph, text, bold=False, size=None):
    """Add a run with Thai font and language settings."""
    run = paragraph.add_run(text)
    run.font.name = THAI_FONT
    run._element.rPr.rFonts.set(qn('w:cs'), THAI_FONT)
    run.font.bold = bold
    if size:
        run.font.size = size
        rPr = run._element.get_or_add_rPr()
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(size.pt * 2)))
        rPr.append(szCs)
    if bold:
        rPr = run._element.get_or_add_rPr()
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)
    set_thai_language(run)
    return run

def disable_spell_check_paragraph(paragraph):
    """Disable spell checking on the entire paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    noProof = OxmlElement('w:noProof')
    rPr.append(noProof)
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'th-TH')
    lang.set(qn('w:bidi'), 'th-TH')
    rPr.append(lang)

def add_page_break_before(paragraph):
    """Add a page break before this paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    pageBreak = OxmlElement('w:pageBreakBefore')
    pPr.append(pageBreak)

def setup_page(doc):
    """Configure A4 page layout."""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)
    
    # Header with gazette info
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disable_spell_check_paragraph(hp)
    add_thai_run(hp, "ราชกิจจานุเบกษา  ฉบับพิเศษ", size=Pt(12))
    
    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disable_spell_check_paragraph(fp)
    
    # Add page number field
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    run2 = fp.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)
    
    run3 = fp.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)

def finalize_document(doc):
    """Add document-level settings to prevent spell check."""
    settings_part = doc.settings.element
    proofState = OxmlElement('w:proofState')
    proofState.set(qn('w:spelling'), 'clean')
    proofState.set(qn('w:grammar'), 'clean')
    settings_part.append(proofState)
    
    themeFontLang = settings_part.find(qn('w:themeFontLang'))
    if themeFontLang is None:
        themeFontLang = OxmlElement('w:themeFontLang')
        settings_part.append(themeFontLang)
    themeFontLang.set(qn('w:val'), 'th-TH')
    themeFontLang.set(qn('w:bidi'), 'th-TH')

# ==================== Read Processed Text ====================

if not os.path.exists(processed_text_path):
    print(f"Error: {processed_text_path} not found. Run generate_pdf.py first.")
    sys.exit(1)

with open(processed_text_path, 'r', encoding='utf-8') as f:
    text = f.read()

# ==================== Parse Text Structure ====================

idx_chapter_1 = text.find("หมวด ๑")
if idx_chapter_1 == -1:
    print("Error: Could not find 'หมวด ๑' in text.")
    sys.exit(1)

cover_and_preamble = text[:idx_chapter_1].strip()
body_text = text[idx_chapter_1:].strip()

cover_paragraphs = [p.strip() for p in re.split(r'\n\s*\n', cover_and_preamble) if p.strip()]
body_paragraphs = [p.strip() for p in re.split(r'\n\s*\n', body_text) if p.strip()]

# ==================== Create Main Word Document ====================

print("สร้างเอกสาร Word หลัก...")
doc = Document()
setup_page(doc)
set_document_thai_language(doc)

# Configure Normal style
style = doc.styles['Normal']
font = style.font
font.name = THAI_FONT
font.size = Pt(16)
style.element.rPr.rFonts.set(qn('w:cs'), THAI_FONT)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(3)
pf.line_spacing = 1.15

# ==================== Cover Page ====================

# Title
if cover_paragraphs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(18)
    disable_spell_check_paragraph(p)
    add_thai_run(p, cover_paragraphs[0], bold=True, size=Pt(22))

# King section
if len(cover_paragraphs) > 1:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    disable_spell_check_paragraph(p)
    for line in cover_paragraphs[1].split('\n'):
        line = line.strip()
        if line:
            add_thai_run(p, line + '\n', size=Pt(16))

# Preamble
for para_text in cover_paragraphs[2:]:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(2)
    p.paragraph_format.space_after = Pt(6)
    disable_spell_check_paragraph(p)
    add_thai_run(p, para_text, size=Pt(16))

# ==================== Body ====================

art_counter = 1
is_first_chapter = True
toc_entries = []  # (chapter_title, first_article, last_article)
current_chapter_title = ""
current_chapter_first_art = 0
chapter_count = 0

for i, paragraph_text in enumerate(body_paragraphs):
    is_chapter = paragraph_text.startswith("หมวด") or paragraph_text.startswith("=")
    is_special_chapter = paragraph_text.startswith("=") or (paragraph_text.startswith("หมวดพิเศษ"))
    is_transitional = paragraph_text.startswith("บทเฉพาะกาล")
    
    if is_chapter or is_transitional:
        # Save previous chapter's article range
        if current_chapter_title and current_chapter_first_art > 0:
            toc_entries.append((current_chapter_title, current_chapter_first_art, art_counter - 1))
        
        # Chapter title — always start on new page
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        disable_spell_check_paragraph(p)
        
        # Page break before every chapter (except first which follows preamble)
        if not is_first_chapter:
            add_page_break_before(p)
        is_first_chapter = False
        
        # Clean up the title text
        clean_title = paragraph_text.replace("=", "").strip()
        
        for line in clean_title.split('\n'):
            line = line.strip()
            if line:
                add_thai_run(p, line, bold=True, size=Pt(18))
                add_thai_run(p, '\n')
        
        current_chapter_title = clean_title.replace('\n', ' ')
        current_chapter_first_art = art_counter
        chapter_count += 1
                
    elif paragraph_text.startswith("ส่วน"):
        # Part title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        disable_spell_check_paragraph(p)
        for line in paragraph_text.split('\n'):
            line = line.strip()
            if line:
                add_thai_run(p, line, bold=True, size=Pt(16))
                add_thai_run(p, '\n')
                
    elif paragraph_text.startswith("มาตรา"):
        # Article — renumber sequentially
        match = re.match(r'^(มาตรา\s+[๐-๙/\s]+(?:ทวิ|ตรี|จัตวา)?)\s+(.*)', paragraph_text, re.DOTALL)
        if match:
            _, content = match.groups()
            new_art_label = f"มาตรา {to_thai_numeral(art_counter)}"
            art_counter += 1
            
            lines = content.split('\n')
            
            # First line
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(3)
            disable_spell_check_paragraph(p)
            add_thai_run(p, new_art_label + "  ", bold=True, size=Pt(16))
            if lines:
                add_thai_run(p, lines[0].strip(), size=Pt(16))
            
            # Subsequent lines
            for line in lines[1:]:
                line = line.strip()
                if not line:
                    continue
                list_match = re.match(r'^(\([๐-๙]+\))\s*(.*)', line)
                if list_match:
                    num, item_text = list_match.groups()
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(2.2)
                    p.paragraph_format.first_line_indent = Cm(-0.7)
                    p.paragraph_format.space_after = Pt(1)
                    disable_spell_check_paragraph(p)
                    add_thai_run(p, f"{num} ", size=Pt(16))
                    add_thai_run(p, item_text, size=Pt(16))
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Cm(1.5)
                    p.paragraph_format.space_after = Pt(3)
                    disable_spell_check_paragraph(p)
                    add_thai_run(p, line, size=Pt(16))
        else:
            p = doc.add_paragraph()
            disable_spell_check_paragraph(p)
            add_thai_run(p, paragraph_text, size=Pt(16))
    
    elif paragraph_text.startswith("ผู้รับสนอง"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(24)
        disable_spell_check_paragraph(p)
        add_thai_run(p, "ผู้รับสนองพระบรมราชโองการ", size=Pt(16))
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        disable_spell_check_paragraph(p)
        add_thai_run(p, "นายกรัฐมนตรี", bold=True, size=Pt(16))
    
    else:
        # Normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.5)
        p.paragraph_format.space_after = Pt(3)
        disable_spell_check_paragraph(p)
        for line in paragraph_text.split('\n'):
            line = line.strip()
            if line:
                add_thai_run(p, line, size=Pt(16))
                add_thai_run(p, '\n')

# Save last chapter entry
if current_chapter_title and current_chapter_first_art > 0:
    toc_entries.append((current_chapter_title, current_chapter_first_art, art_counter - 1))

finalize_document(doc)
doc.save(word_output_path)
file_size = os.path.getsize(word_output_path)
total_articles = art_counter - 1
print(f"✅ เอกสาร Word สร้างเรียบร้อย!")
print(f"   Path: {word_output_path}")
print(f"   Size: {file_size:,} bytes")
print(f"   จำนวนมาตรา: {total_articles}")
print(f"   จำนวนหมวด: {chapter_count}")

# ==================== Create TOC Document ====================

print("\nสร้างไฟล์สารบัญ...")
toc_doc = Document()
setup_page(toc_doc)
set_document_thai_language(toc_doc)

# Style
toc_style = toc_doc.styles['Normal']
toc_style.font.name = THAI_FONT
toc_style.font.size = Pt(16)
toc_style.element.rPr.rFonts.set(qn('w:cs'), THAI_FONT)

# Title
p = toc_doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
p.paragraph_format.space_after = Pt(6)
disable_spell_check_paragraph(p)
add_thai_run(p, "สารบัญ", bold=True, size=Pt(24))

p = toc_doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
disable_spell_check_paragraph(p)
add_thai_run(p, "รัฐธรรมนูญแห่งราชอาณาจักรไทย พุทธศักราช ๒๕๒๕", bold=True, size=Pt(18))

# Horizontal line
p = toc_doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
disable_spell_check_paragraph(p)
add_thai_run(p, "─" * 60, size=Pt(10))

# Table header
table = toc_doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

# Set column widths
for cell in table.columns[0].cells:
    cell.width = Cm(10)
for cell in table.columns[1].cells:
    cell.width = Cm(3)
for cell in table.columns[2].cells:
    cell.width = Cm(3)

# Header row
hdr = table.rows[0]
for cell in hdr.cells:
    # Grey background
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '2C3E50')
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)

h0 = hdr.cells[0].paragraphs[0]
h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
disable_spell_check_paragraph(h0)
add_thai_run(h0, "หมวด / รายการ", bold=True, size=Pt(16)).font.color.rgb = RGBColor(255, 255, 255)

h1 = hdr.cells[1].paragraphs[0]
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
disable_spell_check_paragraph(h1)
add_thai_run(h1, "มาตราเริ่มต้น", bold=True, size=Pt(14)).font.color.rgb = RGBColor(255, 255, 255)

h2 = hdr.cells[2].paragraphs[0]
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
disable_spell_check_paragraph(h2)
add_thai_run(h2, "มาตราสุดท้าย", bold=True, size=Pt(14)).font.color.rgb = RGBColor(255, 255, 255)

# Data rows
for idx, (title, first_art, last_art) in enumerate(toc_entries):
    row = table.add_row()
    
    # Alternate row coloring
    if idx % 2 == 0:
        bg_color = 'EBF5FB'
    else:
        bg_color = 'FFFFFF'
    
    for cell in row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), bg_color)
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)
    
    c0 = row.cells[0].paragraphs[0]
    disable_spell_check_paragraph(c0)
    add_thai_run(c0, title, bold=False, size=Pt(14))
    
    c1 = row.cells[1].paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disable_spell_check_paragraph(c1)
    add_thai_run(c1, to_thai_numeral(first_art), size=Pt(14))
    
    c2 = row.cells[2].paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disable_spell_check_paragraph(c2)
    add_thai_run(c2, to_thai_numeral(last_art), size=Pt(14))

# Summary footer
p = toc_doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
disable_spell_check_paragraph(p)
add_thai_run(p, f"รวมทั้งสิ้น {to_thai_numeral(chapter_count)} หมวด  {to_thai_numeral(total_articles)} มาตรา", bold=True, size=Pt(16))

finalize_document(toc_doc)
toc_doc.save(toc_output_path)
toc_size = os.path.getsize(toc_output_path)
print(f"✅ สารบัญสร้างเรียบร้อย!")
print(f"   Path: {toc_output_path}")
print(f"   Size: {toc_size:,} bytes")
print(f"   จำนวนรายการ: {len(toc_entries)}")
