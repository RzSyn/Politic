# generate_early_word.py — สร้างไฟล์ Word (.docx) รัฐธรรมนูญฉบับปฐมเวลา ๒๓๗๕ และสารบัญแยกไฟล์
# ตั้งค่าภาษาไทย + ฟอนต์ TH Sarabun New ไม่ให้ขึ้นเส้นแดง (spell check)

import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

wd = r"c:\Users\ASUA\OneDrive\Documents\รัฐธรรมนุญจำลอง"
output_doc_path = os.path.join(wd, "รัฐธรรมนูญ_ฉบับปฐมเวลา_๒๓๗๕.docx")
output_toc_path = os.path.join(wd, "สารบัญ_รัฐธรรมนูญ_ฉบับปฐมเวลา_๒๓๗๕.docx")

# Font settings
THAI_FONT = "TH Sarabun New"

def set_thai_language(run):
    """Set the proofing language of a run to Thai (th-TH)."""
    rPr = run._element.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'th-TH')
    lang.set(qn('w:bidi'), 'th-TH')
    lang.set(qn('w:eastAsia'), 'th-TH')
    rPr.append(lang)

def set_document_thai_language(doc):
    """Set the entire document's default language to Thai."""
    style = doc.styles['Normal']
    rPr = style.element.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'th-TH')
    lang.set(qn('w:bidi'), 'th-TH')
    lang.set(qn('w:eastAsia'), 'th-TH')
    rPr.append(lang)
    
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = OxmlElement('w:docDefaults')
        styles_element.insert(0, docDefaults)
    
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault')
        docDefaults.append(rPrDefault)
    
    rPr2 = rPrDefault.find(qn('w:rPr'))
    if rPr2 is None:
        rPr2 = OxmlElement('w:rPr')
        rPrDefault.append(rPr2)
    
    lang2 = OxmlElement('w:lang')
    lang2.set(qn('w:val'), 'th-TH')
    lang2.set(qn('w:bidi'), 'th-TH')
    lang2.set(qn('w:eastAsia'), 'th-TH')
    rPr2.append(lang2)
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), THAI_FONT)
    rFonts.set(qn('w:hAnsi'), THAI_FONT)
    rFonts.set(qn('w:cs'), THAI_FONT)
    rFonts.set(qn('w:eastAsia'), THAI_FONT)
    rPr2.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '32')
    rPr2.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '32')
    rPr2.append(szCs)

def add_thai_run(paragraph, text, bold=False, size=None):
    """Add a run with Thai font and language settings."""
    run = paragraph.add_run(text)
    run.font.name = THAI_FONT
    run._element.rPr.rFonts.set(qn('w:cs'), THAI_FONT)
    run.font.bold = bold
    if size:
        run.font.size = size
        rPr = run._element.get_or_add_rPr()
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(size.pt * 2)))
        rPr.append(szCs)
    if bold:
        rPr = run._element.get_or_add_rPr()
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)
    set_thai_language(run)
    return run

def disable_spell_check_paragraph(paragraph):
    """Disable spell checking on the entire paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    noProof = OxmlElement('w:noProof')
    rPr.append(noProof)
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'th-TH')
    lang.set(qn('w:bidi'), 'th-TH')
    rPr.append(lang)

def add_page_break_before(paragraph):
    """Add a page break before this paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    pageBreak = OxmlElement('w:pageBreakBefore')
    pPr.append(pageBreak)

def setup_page(doc):
    """Configure A4 page layout."""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)
    
    # Header with gazette info
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disable_spell_check_paragraph(hp)
    add_thai_run(hp, "ราชกิจจานุเบกษา  ฉบับพิเศษ", size=Pt(12))
    
    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disable_spell_check_paragraph(fp)
    
    # Add page number field
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    
    run2 = fp.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)
    
    run3 = fp.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)

def finalize_document(doc):
    """Add document-level settings to prevent spell check."""
    settings_part = doc.settings.element
    proofState = OxmlElement('w:proofState')
    proofState.set(qn('w:spelling'), 'clean')
    proofState.set(qn('w:grammar'), 'clean')
    settings_part.append(proofState)
    
    themeFontLang = settings_part.find(qn('w:themeFontLang'))
    if themeFontLang is None:
        themeFontLang = OxmlElement('w:themeFontLang')
        settings_part.append(themeFontLang)
    themeFontLang.set(qn('w:val'), 'th-TH')
    themeFontLang.set(qn('w:bidi'), 'th-TH')

# Early constitution data structure
early_constitution = [
    {
        "title": "หมวดที่ ๑ บททั่วไปและพระราชอำนาจ (มาตรา ๑ - ๑๐)",
        "first_art": 1,
        "last_art": 10,
        "articles": [
            ("มาตรา ๑", "รัฐธรรมนูญนี้มีชื่อว่า \"รัฐธรรมนูญแห่งราชอาณาจักรไทย ฉบับปฐมเวลา\""),
            ("มาตรา ๒", "ประเทศไทยเป็นราชอาณาจักรอันหนึ่งอันเดียว จะแบ่งแยกมิได้"),
            ("มาตรา ๓", "อำนาจอธิปไตยเป็นของชาติ เพื่อให้เกิดความมั่นคงในระยะเปลี่ยนผ่าน การใช้อำนาจจะถูกแบ่งสัดส่วนโดยพระมหากษัตริย์ทรงใช้อำนาจบริหารและวินิจฉัยสูงสุด (ร้อยละ ๖๐) และราษฎรใช้อำนาจผ่านการเลือกตั้งระดับท้องถิ่นและการให้คำปรึกษา (ร้อยละ ๔๐)"),
            ("มาตรา ๔", "พระมหากษัตริย์ทรงเป็นประมุขแห่งรัฐ ทรงดำรงอยู่ในฐานะอันเป็นที่เคารพสักการะ ผู้ใดจะละเมิดมิได้"),
            ("มาตรา ๕", "พระมหากษัตริย์ทรงใช้อำนาจบริหารผ่านกลไกของรัฐตามที่ระบุไว้ในรัฐธรรมนูญนี้"),
            ("มาตรา ๖", "พระมหากษัตริย์ทรงใช้อำนาจนิติบัญญัติร่วมกับสภาที่ปรึกษาราษฎร"),
            ("มาตรา ๗", "พระมหากษัตริย์ทรงใช้อำนาจตุลาการผ่านศาลยุติธรรมที่ทรงแต่งตั้ง"),
            ("มาตรา ๘", "พระมหากษัตริย์ทรงดำรงตำแหน่งจอมทัพไทย"),
            ("มาตรา ๙", "พระมหากษัตริย์ทรงไว้ซึ่งพระราชอำนาจเด็ดขาด (Veto) ในการยับยั้งกฎหมายหรือมติใดๆ ของสภาที่ทรงเห็นว่าจะเป็นภัยต่อความมั่นคงของแผ่นดิน"),
            ("มาตรา ๑๐", "การสืบราชสมบัติให้เป็นไปตามกฎมณเฑียรบาลและพระราชวินิจฉัยของพระมหากษัตริย์")
        ]
    },
    {
        "title": "หมวดที่ ๒ สิทธิ หน้าที่ และส่วนร่วมของราษฎร (มาตรา ๑๑ - ๒๐)",
        "first_art": 11,
        "last_art": 20,
        "articles": [
            ("มาตรา ๑๑", "บุคคลย่อมเสมอกันในกฎหมาย มีสิทธิและเสรีภาพขั้นพื้นฐานเท่าเทียมกัน"),
            ("มาตรา ๑๒", "ราษฎรมีเสรีภาพในการแสดงความคิดเห็น ตราบเท่าที่ไม่ขัดต่อความสงบเรียบร้อยและพระราชอำนาจ"),
            ("มาตรา ๑๓", "บุคคลมีเสรีภาพในการนับถือศาสนา"),
            ("มาตรา ๑๔", "บุคคลมีสิทธิในทรัพย์สินและการสืบมรดกของตน"),
            ("มาตรา ๑๕", "ราษฎรมีสิทธิในการถวายฎีการ้องทุกข์ต่อพระมหากษัตริย์โดยตรงผ่านกลไกที่รัฐกำหนด"),
            ("มาตรา ๑๖", "บุคคลมีหน้าที่ป้องกันประเทศและรักษาความมั่นคงของชาติ"),
            ("มาตรา ๑๗", "บุคคลมีหน้าที่เสียภาษีอากรตามที่กฎหมายบัญญัติ"),
            ("มาตรา ๑๘", "บุคคลมีหน้าที่เคารพกฎหมายและปฏิบัติตามรัฐธรรมนูญอย่างเคร่งครัด"),
            ("มาตรา ๑๙", "ราษฎรในแต่ละท้องถิ่นมีสิทธิออกเสียงเลือกตั้งผู้แทนในระดับหมู่บ้านและตำบลของตน (อันเป็นจุดเริ่มต้นของสัดส่วนอำนาจร้อยละ ๔๐)"),
            ("มาตรา ๒๐", "รัฐมีหน้าที่ต้องให้การศึกษาด้านการปกครองแก่ราษฎร เพื่อเตรียมความพร้อมสู่การมีส่วนร่วมในระดับชาติ")
        ]
    },
    {
        "title": "หมวดที่ ๓ การปกครองส่วนท้องถิ่น (มาตรา ๒๑ - ๓๐)",
        "first_art": 21,
        "last_art": 30,
        "articles": [
            ("มาตรา ๒๑", "การปกครองส่วนท้องถิ่นคือรากฐานสำคัญ รัฐต้องส่งเสริมให้ราษฎรเรียนรู้การปกครองตนเองจากสเกลที่เล็กที่สุดก่อน"),
            ("มาตรา ๒๒", "ให้มีการจัดตั้ง \"สภาหมู่บ้าน\" ขึ้นในทุกชุมชน"),
            ("มาตรา ๒๓", "ราษฎรมีสิทธิเลือกตั้งผู้ใหญ่บ้านและสมาชิกสภาหมู่บ้านโดยตรง"),
            ("มาตรา ๒๔", "ให้สภาหมู่บ้านคัดเลือกตัวแทนเพื่อประกอบกันเป็น \"สภาตำบล\""),
            ("มาตรา ๒๕", "การเลือกตั้งระดับเมืองและมณฑลจะค่อยๆ ถูกขยายผลขึ้นเมื่อสภาตำบลมีความเข้มแข็งตามเกณฑ์ที่รัฐกำหนด"),
            ("มาตรา ๒๖", "สภาท้องถิ่นมีอำนาจในการจัดการภาษีท้องถิ่นร้อยละ ๔๐ เพื่อบำรุงรักษาพื้นที่ของตน"),
            ("มาตรา ๒๗", "พระมหากษัตริย์จะทรงแต่งตั้งข้าหลวงไปกำกับดูแลการทำงานของสภาท้องถิ่น เพื่อประคับประคองการบริหาร"),
            ("มาตรา ๒๘", "หากสภาท้องถิ่นใดเกิดความทุจริต ข้าหลวงมีอำนาจยุบสภานั้นเพื่อจัดตั้งการเลือกตั้งใหม่"),
            ("มาตรา ๒๙", "กระบวนการเลือกตั้งท้องถิ่นนี้จะถือเป็นเวทีทดลอง เพื่อให้ราษฎรคุ้นเคยกับระบบตัวแทน"),
            ("มาตรา ๓๐", "การใช้อำนาจของสภาท้องถิ่นถือเป็นการสะท้อนเจตนารมณ์ของประชาชนตามสัดส่วนอำนาจที่ได้รับมอบหมาย")
        ]
    },
    {
        "title": "หมวดที่ ๔ นิติบัญญัติและคณะบริหารส่วนกลาง (มาตรา ๓๑ - ๔๐)",
        "first_art": 31,
        "last_art": 40,
        "articles": [
            ("มาตรา ๓๑", "ในส่วนกลาง ให้มี \"สภาที่ปรึกษาราษฎร\" ทำหน้าที่ถวายคำแนะนำแด่พระมหากษัตริย์"),
            ("มาตรา ๓๒", "สภาที่ปรึกษาราษฎรประกอบด้วยผู้ทรงคุณวุฒิที่พระมหากษัตริย์ทรงแต่งตั้ง และตัวแทนที่ได้รับเลือกจากสภาท้องถิ่น"),
            ("มาตรา ๓๓", "สภาที่ปรึกษาฯ มีอำนาจเสนอร่างกฎหมาย แต่การอนุมัติบังคับใช้เป็นพระราชอำนาจของพระมหากษัตริย์"),
            ("มาตรา ๓๔", "รัฐมนตรีแต่ละกระทรวงรับผิดชอบการบริหารราชการแผ่นดินโดยตรงต่อพระมหากษัตริย์"),
            ("มาตรา ๓๕", "สภาที่ปรึกษาฯ สามารถตั้งกระทู้ถามรัฐมนตรีได้ แต่ไม่มีอำนาจลงมติไม่ไว้วางใจ"),
            ("มาตรา ๓๖", "ศาลยุติธรรมมีอิสระในการพิจารณาพิพากษาอรรถคดีในพระปรมาภิไธยพระมหากษัตริย์"),
            ("มาตรา ๓๗", "การทำสนธิสัญญาหรือประกาศสงคราม เป็นพระราชอำนาจเด็ดขาดของพระมหากษัตริย์"),
            ("มาตรา ๓๘", "การแก้ไขอัตราภาษีระดับชาติ ต้องได้รับความเห็นชอบจากสภาที่ปรึกษาฯ ก่อนนำขึ้นทูลเกล้าฯ"),
            ("มาตรา ๓๙", "ในยามฉุกเฉิน พระมหากษัตริย์ทรงไว้ซึ่งอำนาจในการออกพระราชกำหนดที่มีผลบังคับใช้ดังพระราชบัญญัติ"),
            ("มาตรา ๔๐", "กลไกการบริหารทั้งหมดนี้ให้อยู่ภายใต้การกำกับของพระมหากษัตริย์จนกว่าจะสิ้นสุดระยะเวลาเตรียมความพร้อม")
        ]
    },
    {
        "title": "หมวดที่ ๕ บทเฉพาะกาล (มาตรา ๔๑ - ๕๐)",
        "first_art": 41,
        "last_art": 50,
        "articles": [
            ("มาตรา ๔๑", "รัฐธรรมนูญฉบับปฐมเวลานี้ ให้มีผลบังคับใช้ตั้งแต่วันที่ ๓๑ ธันวาคม พุทธศักราช ๒๓๗๕ เป็นต้นไป"),
            ("มาตรา ๔๒", "ให้กำหนดระยะเวลาเตรียมความพร้อมของแผ่นดิน (Transition Period) เป็นเวลาทั้งสิ้น ๔๖ ปี นับแต่วันประกาศใช้รัฐธรรมนูญนี้"),
            ("มาตรา ๔๓", "ตลอดระยะเวลา ๔๖ ปีนี้ จะยังไม่มีตำแหน่งนายกรัฐมนตรี พระมหากษัตริย์จะทรงเป็นผู้นำฝ่ายบริหารสูงสุดด้วยพระองค์เอง"),
            ("มาตรา ๔๔", "ในช่วงเวลาดังกล่าว รัฐต้องเร่งขยายการศึกษาและการเลือกตั้งท้องถิ่นให้ครอบคลุมทุกมณฑลทั่วราชอาณาจักร"),
            ("มาตรา ๔๕", "หากพบว่าท้องถิ่นใดยังไม่พร้อม รัฐสามารถขยายเวลาการริเริ่มการเลือกตั้งในพื้นที่นั้นออกไปได้"),
            ("มาตรา ๔๖", "เมื่อระยะเวลาเตรียมความพร้อม ๔๖ ปี สิ้นสุดลง (พุทธศักราช ๒๔๒๑) ให้ถือว่าการเรียนรู้ของราษฎรบรรลุผลในขั้นปฐมภูมิ"),
            ("มาตรา ๔๗", "หลังผ่านพ้น ๔๖ ปี พระมหากษัตริย์จะทรงใช้พระราชอำนาจในการโปรดเกล้าฯ แต่งตั้ง \"นายกรัฐมนตรีคนแรก\" ของราชอาณาจักร"),
            ("มาตรา ๔๘", "นายกรัฐมนตรีคนแรกจะมีหน้าที่จัดตั้งคณะรัฐมนตรี เพื่อรับมอบอำนาจบริหารร้อยละ ๖๐ จากพระมหากษัตริย์มาดำเนินการในรูปแบบรัฐบาล"),
            ("มาตรา ๔๙", "ภายหลังจากการมีนายกรัฐมนตรีคนแรกแล้ว ให้พระมหากษัตริย์และราษฎร (ผ่านสภาตัวแทน) ร่วมกันประเมินผลสัมฤทธิ์ของรัฐธรรมนูญฉบับปฐมเวลานี้"),
            ("มาตรา ๕๐", "การจะร่างรัฐธรรมนูญฉบับใหม่เพื่อเปลี่ยนผ่านสัดส่วนอำนาจไปสู่รูปแบบอื่นหรือไม่นั้น ให้อยู่ที่ดุลพินิจร่วมกันของพระมหากษัตริย์และประชาชน หากเห็นพ้องว่ายังไม่พร้อม รัฐธรรมนูญฉบับปฐมเวลานี้จะยังคงมีผลบังคับใช้สืบไป")
        ]
    }
]

# ==================== Generate Main Early Constitution Document ====================

print("สร้างเอกสาร รัฐธรรมนูญ_ฉบับปฐมเวลา_๒๓๗๕.docx ...")
doc = Document()
setup_page(doc)
set_document_thai_language(doc)

# Configure Normal style
style = doc.styles['Normal']
font = style.font
font.name = THAI_FONT
font.size = Pt(16)
style.element.rPr.rFonts.set(qn('w:cs'), THAI_FONT)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(3)
pf.line_spacing = 1.15

# --- Cover Page ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after = Pt(18)
disable_spell_check_paragraph(p)
add_thai_run(p, "รัฐธรรมนูญแห่งราชอาณาจักรไทย\nฉบับปฐมเวลา", bold=True, size=Pt(22))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
disable_spell_check_paragraph(p)
add_thai_run(p, "ประกาศใช้ ณ วันที่ ๓๑ ธันวาคม พุทธศักราช ๒๓๗๕", bold=True, size=Pt(16))

# --- Preamble (คำปรารภ) ---
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(2)
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(12)
disable_spell_check_paragraph(p)
add_thai_run(p, "คำปรารภ:", bold=True, size=Pt(16))

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(2)
p.paragraph_format.space_after = Pt(12)
disable_spell_check_paragraph(p)
add_thai_run(p, "เพื่อให้ราชอาณาจักรไทยก้าวสู่ความเจริญทัดเทียมนานาอารยประเทศ และเพื่อให้ราษฎรได้เรียนรู้วิถีแห่งการปกครองตนเอง พระมหากษัตริย์จึงมีพระบรมราชโองการโปรดเกล้าฯ พระราชทานรัฐธรรมนูญฉบับปฐมเวลานี้ เพื่อเป็นบันไดขั้นแรกในการกระจายอำนาจสู่ท้องถิ่น และเตรียมความพร้อมของแผ่นดินในระยะเวลา ๔๖ ปีข้างหน้า", size=Pt(16))

# --- Chapters and Articles ---
is_first = True
for ch in early_constitution:
    # Page Break and Chapter Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    disable_spell_check_paragraph(p)
    
    # We do a page break for every chapter (except possibly after the cover/preamble page if desired, 
    # but normally we want to separate chapters on new pages as requested)
    add_page_break_before(p)
    
    add_thai_run(p, ch['title'], bold=True, size=Pt(18))
    
    for art_num, art_content in ch['articles']:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.5)
        p.paragraph_format.space_after = Pt(4)
        disable_spell_check_paragraph(p)
        
        # Format article
        add_thai_run(p, art_num + "  ", bold=True, size=Pt(16))
        add_thai_run(p, art_content, size=Pt(16))

finalize_document(doc)
doc.save(output_doc_path)
doc_size = os.path.getsize(output_doc_path)
print(f"✅ เอกสารรัฐธรรมนูญฉบับปฐมเวลาหลักสร้างเรียบร้อย! Size: {doc_size:,} bytes")


# ==================== Generate Table of Contents Document ====================

print("สร้างเอกสาร สารบัญ_รัฐธรรมนูญ_ฉบับปฐมเวลา_๒๓๗๕.docx ...")
toc_doc = Document()
setup_page(toc_doc)
set_document_thai_language(toc_doc)

# Style
toc_style = toc_doc.styles['Normal']
toc_style.font.name = THAI_FONT
toc_style.font.size = Pt(16)
toc_style.element.rPr.rFonts.set(qn('w:cs'), THAI_FONT)

# Title
p = toc_doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
p.paragraph_format.space_after = Pt(6)
disable_spell_check_paragraph(p)
add_thai_run(p, "สารบัญ", bold=True, size=Pt(24))

p = toc_doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
disable_spell_check_paragraph(p)
add_thai_run(p, "รัฐธรรมนูญแห่งราชอาณาจักรไทย ฉบับปฐมเวลา พุทธศักราช ๒๓๗๕", bold=True, size=Pt(16))

# Horizontal line
p = toc_doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
disable_spell_check_paragraph(p)
add_thai_run(p, "─" * 60, size=Pt(10))

# Table
table = toc_doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

# Set column widths
for cell in table.columns[0].cells:
    cell.width = Cm(10)
for cell in table.columns[1].cells:
    cell.width = Cm(3)
for cell in table.columns[2].cells:
    cell.width = Cm(3)

# Header row
hdr = table.rows[0]
for cell in hdr.cells:
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '2C3E50')
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)

h0 = hdr.cells[0].paragraphs[0]
h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
disable_spell_check_paragraph(h0)
add_thai_run(h0, "หมวด / รายการ", bold=True, size=Pt(16)).font.color.rgb = RGBColor(255, 255, 255)

h1 = hdr.cells[1].paragraphs[0]
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
disable_spell_check_paragraph(h1)
add_thai_run(h1, "มาตราเริ่มต้น", bold=True, size=Pt(14)).font.color.rgb = RGBColor(255, 255, 255)

h2 = hdr.cells[2].paragraphs[0]
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
disable_spell_check_paragraph(h2)
add_thai_run(h2, "มาตราสุดท้าย", bold=True, size=Pt(14)).font.color.rgb = RGBColor(255, 255, 255)

# Convert integer to Thai numerals
def to_thai_numeral(num):
    thai_nums = {'0': '๐', '1': '๑', '2': '๒', '3': '๓', '4': '๔',
                 '5': '๕', '6': '๖', '7': '๗', '8': '๘', '9': '๙'}
    return ''.join(thai_nums.get(d, d) for d in str(num))

# Data rows
for idx, ch in enumerate(early_constitution):
    row = table.add_row()
    
    # Alternate row coloring
    bg_color = 'EBF5FB' if idx % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), bg_color)
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)
    
    title_text = ch['title'].split(" (มาตรา")[0].strip() # Clean up chapter title
    
    c0 = row.cells[0].paragraphs[0]
    disable_spell_check_paragraph(c0)
    add_thai_run(c0, title_text, bold=False, size=Pt(14))
    
    c1 = row.cells[1].paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disable_spell_check_paragraph(c1)
    add_thai_run(c1, to_thai_numeral(ch['first_art']), size=Pt(14))
    
    c2 = row.cells[2].paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disable_spell_check_paragraph(c2)
    add_thai_run(c2, to_thai_numeral(ch['last_art']), size=Pt(14))

# Summary footer
p = toc_doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
disable_spell_check_paragraph(p)
add_thai_run(p, f"รวมทั้งสิ้น ๕ หมวด  ๕๐ มาตรา", bold=True, size=Pt(16))

finalize_document(toc_doc)
toc_doc.save(output_toc_path)
toc_size = os.path.getsize(output_toc_path)
print(f"✅ เอกสารสารบัญฉบับปฐมเวลาสร้างเรียบร้อย! Size: {toc_size:,} bytes")
